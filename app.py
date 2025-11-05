import whisper
from flask import Flask, render_template, request, send_file, redirect, url_for
import os, json, datetime
from utils.extract_rules import extract_from_brat
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from PyPDF2 import PdfReader
import speech_recognition as sr
import imageio_ffmpeg as ffmpeg
from pydub import AudioSegment
AudioSegment.converter = r"C:\ffmpeg\ffmpeg\bin\ffmpeg.exe"
AudioSegment.ffmpeg = r"C:\ffmpeg\ffmpeg\bin\ffmpeg.exe"
AudioSegment.ffprobe = r"C:\ffmpeg\ffmpeg\bin\ffprobe.exe"
app = Flask(__name__)
BASE = os.path.dirname(__file__)
UPLOAD_DIR = os.path.join(BASE, "uploads")
EXPORT_DIR = os.path.join(BASE, "exports")
HISTORY_DIR = os.path.join(BASE, "history")

os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(EXPORT_DIR, exist_ok=True)
os.makedirs(HISTORY_DIR, exist_ok=True)
HISTORY_FILE = os.path.join(HISTORY_DIR, "history.json")
# ---------- Helper Functions ----------
def extract_text_from_file(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    text = ""
    if ext == ".pdf":
        reader = PdfReader(file_path)
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read()
    else:
        raise ValueError("Unsupported file type")
    return text.strip()

def save_history(data, original_filename):
    now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    history_item = {
        "id": int(datetime.datetime.now().timestamp()),
        "filename": original_filename,
        "summary": data.get("summary"),
        "key_topics": data.get("key_topics"),
        "decisions": data.get("decisions"),
        "actions": data.get("actions"),
        "timestamp": now
    }
    hist_path = os.path.join(HISTORY_DIR, "history.json")
    history = []
    if os.path.exists(hist_path):
        with open(hist_path, "r", encoding="utf-8") as f:
            history = json.load(f)
    history.insert(0, history_item)
    with open(hist_path, "w", encoding="utf-8") as f:
        json.dump(history, f, indent=2, ensure_ascii=False)
    return history_item

def load_history():
    hist_path = os.path.join(HISTORY_DIR, "history.json")
    if os.path.exists(hist_path):
        with open(hist_path, "r", encoding="utf-8") as f:
            return json.load(f)
    return []

def get_history_item(item_id):
    for item in load_history():
        if str(item["id"]) == str(item_id):
            return item
    return None
def convert_audio_to_text_whisper(file_path):
    try:
        model = whisper.load_model("base")  # You can change to "small", "medium", or "large"
        result = model.transcribe(file_path)
        return result["text"].strip()
    except Exception as e:
        return f"(Whisper transcription failed: {e})"

# ---------- Routes ----------
@app.route('/')
def home():
    """Landing page with two buttons"""
    return render_template('home.html')

@app.route('/add_transcript')
def add_transcript():
    """Opens your existing transcript page"""
    return render_template('index.html')

@app.route('/audio')
def audio_page():
    """Open audio upload page"""
    return render_template('audio.html')

@app.route('/process_audio', methods=['POST'])
def process_audio():
    audio_file = request.files.get('audio')
    if not audio_file:
        return render_template('audio.html', error="Please upload an audio file!")

    filename = audio_file.filename
    file_path = os.path.join(UPLOAD_DIR, filename)
    audio_file.save(file_path)

    # ✅ Convert speech to text using Whisper
    text_output = convert_audio_to_text_whisper(file_path)

    # Show the text result on the same page
    return render_template('audio.html', audio_text=text_output)

def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    text_input = request.form.get('text_input', '').strip()
    upload = request.files.get('file')

    text = ""
    filename = ""

    # Case 1: User typed or pasted text
    if text_input:
        text = text_input
        filename = "Typed Transcript"

    # Case 2: User uploaded a file
    elif upload and upload.filename:
        filename = upload.filename
        file_path = os.path.join(UPLOAD_DIR, filename)
        upload.save(file_path)
        try:
            text = extract_text_from_file(file_path)
        except Exception as e:
            return render_template('index.html', error=f"Failed to read file: {str(e)}")

    else:
        return render_template('index.html', error="Please upload a file or type some text!")

    # Process text
    results = extract_from_brat(text, entities=[], relations=[])
    save_history(results, filename)
    return render_template("index.html", results=results, uploaded=True)

@app.route('/history')
def history():
    # Ensure the history file exists
    os.makedirs(os.path.dirname(HISTORY_FILE), exist_ok=True)
    if not os.path.exists(HISTORY_FILE):
        with open(HISTORY_FILE, 'w', encoding='utf-8') as f:
            f.write("[]")

    # Load safely
    try:
        with open(HISTORY_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        data = []

    # ✅ Don’t filter by 'formatted_summary'; use what you actually saved
    clean_data = []
    for item in data:
        if isinstance(item, dict) and 'summary' in item:
            clean_data.append(item)

    # ✅ Pass cleaned data to template
    return render_template('history.html', history=clean_data)

# ---------- Export Routes ----------
@app.route('/export/docx', methods=['POST'])
def export_docx():
    payload = request.form['payload']
    data = json.loads(payload)

    doc = Document()
    doc.add_heading('AIMS - Meeting Summary', level=1)
    doc.add_paragraph(data.get('summary', ''))

    doc.add_heading('Key Topics', level=2)
    doc.add_paragraph(", ".join(data.get('key_topics', [])))

    doc.add_heading('Decisions', level=2)
    for d in data.get('decisions', []):
        doc.add_paragraph(d, style='List Bullet')

    doc.add_heading('Action Items', level=2)
    actions = data.get('actions', [])
    if isinstance(actions, dict):  # just in case payload was converted differently
        actions = list(actions.values())

    for i, a in enumerate(actions, start=1):
        line = f"Task {i}: {a.get('task','')} | Person: {a.get('person','')} | Due: {a.get('due','')}"
        doc.add_paragraph(line, style='List Number')

    fname = f"summary_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    outp = os.path.join(EXPORT_DIR, fname)
    doc.save(outp)
    return send_file(outp, as_attachment=True)

@app.route('/export/pdf', methods=['POST'])
def export_pdf():
    payload = request.form['payload']
    data = json.loads(payload)

    fname = f"summary_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
    outp = os.path.join(EXPORT_DIR, fname)

    doc = SimpleDocTemplate(outp)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("AIMS - Meeting Summary", styles['Title']))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Summary:", styles['Heading2']))
    story.append(Paragraph(data.get('summary', ''), styles['Normal']))

    story.append(Paragraph("Key Topics:", styles['Heading2']))
    story.append(Paragraph(", ".join(data.get('key_topics', [])), styles['Normal']))

    story.append(Paragraph("Decisions:", styles['Heading2']))
    for d in data.get('decisions', []):
        story.append(Paragraph(d, styles['Normal']))

    story.append(Paragraph("Action Items:", styles['Heading2']))
    actions = data.get('actions', [])
    if isinstance(actions, dict):
        actions = list(actions.values())

    for i, a in enumerate(actions, start=1):
        line = f"Task {i}: {a.get('task','')} | Person: {a.get('person','')} | Due: {a.get('due','')}"
        story.append(Paragraph(line, styles['Normal']))

    doc.build(story)
    return send_file(outp, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
